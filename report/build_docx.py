#!/usr/bin/env python
"""Build report/build/report_draft.docx from the four Turkish drafts.

This is a MEASUREMENT INSTRUMENT, not the submission. It exists to answer one
question: how many paginated pages does the current body text occupy inside the
KYS template's page geometry? There is deliberately no cover page, no table of
contents and no bibliography.

Design rule, in order of precedence:

1. The committed KYS template is the base document. Its styles and its sectPr
   (A4, 2.5 cm margins, header reference) are inherited, never re-declared in
   code. The only formatting this script declares is (a) the three style
   redefinitions in `redefine_styles`, which are written into the style
   definitions themselves rather than onto individual paragraphs, and (b) the
   handful of direct paragraph properties listed in DIRECT_FORMATTING below,
   each of which has a recorded reason.

2. Nothing personal survives. The template's docProps carry a name in
   cp:lastModifiedBy and have no dc:creator element at all, which some tools
   repopulate from the OS user on save. Every core property is therefore set to
   an explicit empty string so the elements exist and are blank.

3. No content is silently dropped. Anything the converter cannot represent is
   appended to `warnings` and printed at the end of the run.

Run:  .venv/Scripts/python.exe report/build_docx.py
"""

from __future__ import annotations

import re
import sys
import hashlib
from pathlib import Path

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

# --------------------------------------------------------------------------
# configuration
# --------------------------------------------------------------------------

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "NSosyal_Inovasyon_2026_-_Proje_Teknik_Raporu_1_eDrmR.docx"
OUT_DIR = ROOT / "report" / "build"
OUT = OUT_DIR / "report_draft.docx"

# Numeric order. 03 does not exist; the gap is in the source material, not here.
SOURCES = [
    ROOT / "report" / "01_veri_ve_deney_kurgusu.md",
    ROOT / "report" / "02_yontem.md",
    ROOT / "report" / "04_bulgular.md",
    ROOT / "report" / "05_sinirliliklar.md",
]

# The one table style that (a) exists in the template and (b) is not one of the
# 18 anonymous auto-generated ones. Asserted at runtime against the template's
# own style list so a typo cannot silently auto-create a new style.
TABLE_STYLE = "TableNormal"

LINE_SPACING = 1.15
QUOTE_INDENT = Cm(1.0)
LIST_INDENT = Cm(0.75)

# Every direct (non-style) paragraph property this script sets, and why.
DIRECT_FORMATTING = [
    ("body/quote/list paragraphs", "alignment = JUSTIFY",
     "2c asks for justification on body paragraphs specifically. Putting it on "
     "the Normal style instead would push it onto the heading styles too, which "
     "are basedOn Normal -- and Heading 1 must stay exactly as the template "
     "ships it."),
    ("blockquote paragraphs", "left indent %s cm" % QUOTE_INDENT.cm,
     "2f: the template defines no Quote style and this script must not create "
     "one, so the indent is the only way to mark a quote."),
    ("list paragraphs", "left indent %s cm + literal marker" % LIST_INDENT.cm,
     "The template defines no List Bullet / List Number style. Using "
     "python-docx's built-ins would auto-create styles that are not in the "
     "template."),
    ("table header cells", "bold runs",
     "Markdown marks the header row structurally, not with emphasis; the bold "
     "carries that structure across into a table style that has no header "
     "banding of its own."),
    ("table cell paragraphs", "space_after = 0",
     "docDefaults set 160 twips after every paragraph, which inflates a "
     "36-table document by roughly a line per cell."),
]

MONOSPACE_MARKERS = ("consolas", "courier", "mono", "menlo", "monaco")

warnings: list[str] = []


# --------------------------------------------------------------------------
# markdown -> blocks
# --------------------------------------------------------------------------

RE_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
RE_HRULE = re.compile(r"^(-{3,}|\*{3,}|_{3,})\s*$")
RE_LIST = re.compile(r"^(\s*)([-*+]|\d+[.)])\s+(.*)$")
RE_SEPCELL = re.compile(r"^:?-{2,}:?$")


def is_table_line(line: str) -> bool:
    return line.lstrip().startswith("|")


def is_block_start(line: str) -> bool:
    """True if `line` cannot be a lazy continuation of a paragraph."""
    return (
        not line.strip()
        or is_table_line(line)
        or line.startswith(">")
        or RE_HEADING.match(line) is not None
        or RE_HRULE.match(line) is not None
        or RE_LIST.match(line) is not None
    )


def parse_blocks(text: str) -> list[tuple]:
    """Return a flat list of (kind, payload). Soft-wrapped lines are joined."""
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: list[tuple] = []
    i = 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1
            continue

        m = RE_HEADING.match(line)
        if m:
            blocks.append(("heading", (len(m.group(1)), m.group(2).strip())))
            i += 1
            continue

        if RE_HRULE.match(line):
            blocks.append(("hrule", None))
            i += 1
            continue

        if is_table_line(line):
            rows = []
            while i < len(lines) and is_table_line(lines[i]):
                rows.append(lines[i].strip())
                i += 1
            blocks.append(("table", rows))
            continue

        if line.startswith(">"):
            buf = []
            while i < len(lines) and lines[i].startswith(">"):
                buf.append(re.sub(r"^>\s?", "", lines[i]).strip())
                i += 1
            blocks.append(("quote", " ".join(x for x in buf if x)))
            continue

        m = RE_LIST.match(line)
        if m:
            ordered = m.group(2)[0].isdigit()
            items: list[list[str]] = []
            while i < len(lines):
                mm = RE_LIST.match(lines[i])
                if mm:
                    items.append([mm.group(2), mm.group(3).strip()])
                    i += 1
                elif items and lines[i].strip() and lines[i][:1] in (" ", "\t"):
                    # lazy continuation of the previous item
                    items[-1][1] += " " + lines[i].strip()
                    i += 1
                else:
                    break
            blocks.append(("list", (ordered, items)))
            continue

        buf = []
        while i < len(lines) and not is_block_start(lines[i]):
            buf.append(lines[i].strip())
            i += 1
        blocks.append(("para", " ".join(buf)))

    return blocks


# --------------------------------------------------------------------------
# inline markup -> runs
# --------------------------------------------------------------------------

def inline_runs(text: str, where: str) -> list[tuple[str, bool, bool, bool]]:
    """Split `text` into (content, bold, italic, is_code) runs.

    A single left-to-right scan, so a code span nested inside a bold span keeps
    both flags. `_` is NOT treated as emphasis: the drafts use none, but they do
    use underscores inside file and field names.
    """
    runs: list[tuple[str, bool, bool, bool]] = []
    buf: list[str] = []
    bold = ital = False
    i = 0

    def flush() -> None:
        if buf:
            runs.append(("".join(buf), bold, ital, False))
            buf.clear()

    while i < len(text):
        if text[i] == "`":
            j = text.find("`", i + 1)
            if j == -1:
                warnings.append("unclosed backtick in %s: %r" % (where, text[:80]))
                buf.append(text[i])
                i += 1
                continue
            flush()
            runs.append((text[i + 1:j], bold, ital, True))
            i = j + 1
            continue
        if text.startswith("**", i):
            flush()
            bold = not bold
            i += 2
            continue
        if text[i] == "*":
            flush()
            ital = not ital
            i += 1
            continue
        buf.append(text[i])
        i += 1

    flush()
    if bold or ital:
        warnings.append("unbalanced emphasis in %s: %r" % (where, text[:80]))
    return runs


def split_cells(row: str) -> list[str]:
    """Split a markdown table row on `|`, ignoring pipes inside code spans."""
    row = row.strip()
    if row.startswith("|"):
        row = row[1:]
    if row.endswith("|"):
        row = row[:-1]
    cells, buf, in_code = [], [], False
    for ch in row:
        if ch == "`":
            in_code = not in_code
            buf.append(ch)
        elif ch == "|" and not in_code:
            cells.append("".join(buf).strip())
            buf = []
        else:
            buf.append(ch)
    cells.append("".join(buf).strip())
    return cells


# --------------------------------------------------------------------------
# template surgery
# --------------------------------------------------------------------------

def describe_style(style) -> str:
    """One-line summary read back out of the style's own XML."""
    rpr = style.element.find(qn("w:rPr"))
    fonts = sz = szcs = bold = None
    if rpr is not None:
        rf = rpr.find(qn("w:rFonts"))
        if rf is not None:
            fonts = "/".join(
                "%s=%s" % (a, rf.get(qn("w:" + a)))
                for a in ("ascii", "hAnsi", "eastAsia", "cs")
                if rf.get(qn("w:" + a))
            ) or None
        s = rpr.find(qn("w:sz"))
        sz = s.get(qn("w:val")) if s is not None else None
        s = rpr.find(qn("w:szCs"))
        szcs = s.get(qn("w:val")) if s is not None else None
        b = rpr.find(qn("w:b"))
        bold = ("off" if b.get(qn("w:val")) in ("0", "false") else "on") if b is not None else None
    return "styleId=%-8s rFonts[%s] sz=%s (%s pt) szCs=%s bold=%s" % (
        style.style_id, fonts, sz,
        (int(sz) / 2 if sz else "inherit"), szcs, bold,
    )


def set_style_font(style, name: str, pt: float, bold: bool) -> None:
    """Write font/size/weight into the STYLE definition, not onto paragraphs."""
    rpr = style.element.get_or_add_rPr()

    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = rpr.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rf)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rf.set(qn("w:" + attr), name)

    half = str(int(round(pt * 2)))
    for tag in ("w:sz", "w:szCs"):
        el = rpr.find(qn(tag))
        if el is None:
            el = rpr.makeelement(qn(tag), {})
            rpr.append(el)
        el.set(qn("w:val"), half)

    b = rpr.find(qn("w:b"))
    bcs = rpr.find(qn("w:bCs"))
    if bold:
        for tag, el in (("w:b", b), ("w:bCs", bcs)):
            if el is None:
                el = rpr.makeelement(qn(tag), {})
                rpr.append(el)
            el.set(qn("w:val"), "1")
    else:
        for el in (b, bcs):
            if el is not None:
                rpr.remove(el)


def strip_body(document) -> tuple[int, int, int]:
    """Remove every body child except sectPr. Returns (paras, tables, other)."""
    body = document.element.body
    sect = body.find(qn("w:sectPr"))
    p = t = o = 0
    for child in list(body):
        if child is sect:
            continue
        tag = child.tag.split("}")[-1]
        if tag == "p":
            p += 1
        elif tag == "tbl":
            t += 1
        else:
            o += 1
        body.remove(child)
    return p, t, o


def clear_core_properties(document) -> None:
    cp = document.core_properties
    for field in ("author", "last_modified_by", "title", "subject",
                  "category", "comments", "keywords"):
        setattr(cp, field, "")


# --------------------------------------------------------------------------
# writing
# --------------------------------------------------------------------------

def emit_runs(paragraph, text: str, where: str, force_bold: bool = False,
              plain: bool = False) -> None:
    """Write `text` into `paragraph` as formatted runs.

    `plain=True` strips the markup delimiters but emits no run formatting at
    all. Used for headings: a run-level bold or italic inside a heading is
    direct formatting, and it shows up in Word's Styles pane as "Heading 2 +
    Italic", which is exactly the thing the headings are supposed not to do.
    Two headings in the drafts contain inline code spans, so without this they
    would be the only two headings carrying an override.
    """
    for content, bold, ital, is_code in inline_runs(text, where):
        if not content:
            continue
        run = paragraph.add_run(content)
        if plain:
            continue
        run.bold = True if force_bold else (True if bold else None)
        # Inline code: body font (inherited -- no rFonts written, so it cannot
        # be monospace) plus italic, so the 411 spans stay visually distinct.
        run.italic = True if (ital or is_code) else None


def add_body_paragraph(document, text: str, where: str, indent=None) -> None:
    p = document.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent is not None:
        pf.left_indent = indent
    emit_runs(p, text, where)


def add_table(document, rows: list[str], where: str) -> None:
    grid = [split_cells(r) for r in rows]
    body_rows = [r for r in grid if not all(RE_SEPCELL.match(c or "-") for c in r)]
    if not body_rows:
        warnings.append("table with no content rows in %s" % where)
        return
    ncols = max(len(r) for r in body_rows)
    if len({len(r) for r in body_rows}) > 1:
        warnings.append(
            "ragged table in %s: row widths %s, padded to %d"
            % (where, sorted({len(r) for r in body_rows}), ncols)
        )

    table = document.add_table(rows=len(body_rows), cols=ncols)
    table.style = document.styles[TABLE_STYLE]
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True

    for ri, row in enumerate(body_rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            emit_runs(para, row[ci] if ci < len(row) else "",
                      "%s table r%d c%d" % (where, ri, ci),
                      force_bold=(ri == 0))


def build() -> None:
    print("=" * 74)
    print("TEMPLATE")
    print("=" * 74)
    data = TEMPLATE.read_bytes()
    print("  path   %s" % TEMPLATE)
    print("  size   %d bytes" % len(data))
    print("  sha256 %s" % hashlib.sha256(data).hexdigest())

    document = docx.Document(str(TEMPLATE))

    styles_before = [s.name for s in document.styles]
    sect_before = document.element.body.find(qn("w:sectPr"))
    pgsz_before = dict(sect_before.find(qn("w:pgSz")).attrib)
    pgmar_before = dict(sect_before.find(qn("w:pgMar")).attrib)

    print()
    print("=" * 74)
    print("STRIP BODY  (2a)")
    print("=" * 74)
    p, t, o = strip_body(document)
    print("  removed %d paragraphs, %d tables, %d other body children" % (p, t, o))

    sect_after = document.element.body.find(qn("w:sectPr"))
    print("  sectPr survived        : %s" % (sect_after is not None))
    print("  sectPr is same element : %s" % (sect_after is sect_before))
    print("  pgSz  before/after     : %s / %s"
          % (dict(pgsz_before), dict(sect_after.find(qn("w:pgSz")).attrib)))
    print("  pgMar before/after     : %s / %s"
          % (dict(pgmar_before), dict(sect_after.find(qn("w:pgMar")).attrib)))
    print("  headerReference kept   : %s"
          % (sect_after.find(qn("w:headerReference")) is not None))
    styles_after = [s.name for s in document.styles]
    print("  styles before/after    : %d / %d  (identical: %s)"
          % (len(styles_before), len(styles_after), styles_before == styles_after))
    print("  body children left     : %d" % len(list(document.element.body)))

    print()
    print("=" * 74)
    print("STYLE REDEFINITION  (2d)")
    print("=" * 74)
    h1, h2, h3 = (document.styles["Heading %d" % n] for n in (1, 2, 3))
    print("  BEFORE")
    for nm, st in (("Heading 1", h1), ("Heading 2", h2), ("Heading 3", h3)):
        print("    %-10s %s" % (nm, describe_style(st)))

    # Heading 1 is deliberately NOT touched.
    set_style_font(h2, "Arial Black", 12, bold=False)
    set_style_font(h3, "Arial", 12, bold=True)

    print("  AFTER")
    for nm, st in (("Heading 1", h1), ("Heading 2", h2), ("Heading 3", h3)):
        print("    %-10s %s" % (nm, describe_style(st)))
    print("  Heading 1 unchanged    : %s" % (describe_style(h1) ==
          "styleId=Balk1    rFonts[ascii=Arial Black/hAnsi=Arial Black/"
          "eastAsia=Arial Black/cs=Arial Black] sz=28 (14.0 pt) szCs=28 bold=None"))

    # 2c: assert 1.15 rather than inheriting docDefaults' 360 twips (1.5 lines).
    # Written into the Normal style, so every paragraph in the document -- body,
    # quote, list, table cell -- inherits it from a style definition.
    normal = document.styles["Normal"]
    normal.paragraph_format.line_spacing = LINE_SPACING
    print()
    print("  Normal line_spacing set to %s (docDefaults said w:line=360 = 1.5)"
          % normal.paragraph_format.line_spacing)

    print()
    print("=" * 74)
    print("CORE PROPERTIES  (2b)")
    print("=" * 74)
    cp = document.core_properties
    print("  before: author=%r last_modified_by=%r title=%r"
          % (cp.author, cp.last_modified_by, cp.title))
    clear_core_properties(document)
    print("  after : author=%r last_modified_by=%r title=%r subject=%r "
          "category=%r comments=%r keywords=%r"
          % (cp.author, cp.last_modified_by, cp.title, cp.subject,
             cp.category, cp.comments, cp.keywords))

    print()
    print("=" * 74)
    print("TABLE STYLE  (2g)")
    print("=" * 74)
    present = [s.name for s in document.styles if s.name == TABLE_STYLE]
    print("  requested          : %r" % TABLE_STYLE)
    print("  present in template: %s" % bool(present))
    if not present:
        raise SystemExit(
            "ABORT: table style %r is not in the template; using it would "
            "auto-create a style, which 2g forbids." % TABLE_STYLE
        )

    print()
    print("=" * 74)
    print("CONTENT")
    print("=" * 74)
    counts = {"heading": 0, "para": 0, "table": 0, "quote": 0, "list": 0, "hrule": 0}
    per_file = []
    prev_kind = None
    separators = 0

    for src in SOURCES:
        if not src.exists():
            raise SystemExit("ABORT: missing source %s" % src)
        blocks = parse_blocks(src.read_text(encoding="utf-8"))
        local = dict.fromkeys(counts, 0)

        for kind, payload in blocks:
            local[kind] = local.get(kind, 0) + 1
            counts[kind] += 1
            where = src.name

            if kind == "heading":
                level, text = payload
                if level > 3:
                    warnings.append("H%d in %s flattened to Heading 3: %r"
                                    % (level, where, text))
                    level = 3
                para = document.add_paragraph(style="Heading %d" % level)
                emit_runs(para, text, where, plain=True)

            elif kind == "para":
                add_body_paragraph(document, payload, where)

            elif kind == "quote":
                add_body_paragraph(document, payload, where, indent=QUOTE_INDENT)

            elif kind == "list":
                ordered, items = payload
                for n, (marker, text) in enumerate(items, start=1):
                    bullet = ("%d. " % n) if ordered else "• "
                    add_body_paragraph(document, bullet + text, where,
                                       indent=LIST_INDENT)

            elif kind == "table":
                # Two <w:tbl> siblings with no <w:p> between them are merged
                # into a single table by Word -- the second table's columns get
                # forced into the first one's grid. The source has one such
                # pair. An empty separator paragraph is the standard OOXML fix;
                # it is structural, not invented content.
                if prev_kind == "table":
                    document.add_paragraph()
                    separators += 1
                    warnings.append(
                        "%s: two tables are adjacent in the source; inserted an "
                        "empty separator paragraph so Word does not merge them"
                        % where
                    )
                add_table(document, payload, where)

            elif kind == "hrule":
                # A thematic break carries no text. Rendered as an empty
                # paragraph so the visual separation survives without inventing
                # a border style.
                document.add_paragraph()

            prev_kind = kind

        per_file.append((src.name, local, len(blocks)))

    for name, local, total in per_file:
        print("  %-32s blocks=%-4d %s" % (name, total, local))
    print("  %-32s %s" % ("TOTAL", counts))
    print("  %-32s %d" % ("table separator paragraphs", separators))

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(str(OUT))

    print()
    print("=" * 74)
    print("VERIFY OUTPUT  (read back from the saved file)")
    print("=" * 74)
    verify()

    print()
    print("=" * 74)
    print("DIRECT (NON-STYLE) FORMATTING THIS BUILD EMITS")
    print("=" * 74)
    for target, what, why in DIRECT_FORMATTING:
        print("  %s -- %s" % (target, what))
        print("      %s" % why)

    print()
    print("=" * 74)
    print("WARNINGS  (%d)" % len(warnings))
    print("=" * 74)
    for w in warnings:
        print("  ! %s" % w)
    if not warnings:
        print("  none")


def verify() -> None:
    import zipfile
    from lxml import etree

    out = docx.Document(str(OUT))
    print("  output       %s (%d bytes)" % (OUT, OUT.stat().st_size))
    print("  paragraphs   %d" % len(out.paragraphs))
    print("  tables       %d" % len(out.tables))

    kids = [el.tag.split("}")[-1] for el in out.element.body.iterchildren()]
    adjacent = sum(1 for i in range(len(kids) - 1)
                   if kids[i] == "tbl" and kids[i + 1] == "tbl")
    print("  adjacent tbl->tbl pairs (Word merges these): %d  (must be 0)" % adjacent)
    if adjacent:
        raise SystemExit("ABORT: %d adjacent table pairs would be merged by Word"
                         % adjacent)

    heads = {}
    direct_on_headings = 0
    for para in out.paragraphs:
        if para.style.name.startswith("Heading"):
            heads[para.style.name] = heads.get(para.style.name, 0) + 1
            if para.paragraph_format.element.pPr is not None:
                pPr = para.paragraph_format.element.pPr
                if len(pPr):
                    extra = [c.tag.split("}")[-1] for c in pPr
                             if c.tag.split("}")[-1] != "pStyle"]
                    if extra:
                        direct_on_headings += 1
            for run in para.runs:
                if run.font.name or run.font.size or run.font.bold:
                    direct_on_headings += 1
                    break
    print("  headings     %s" % heads)
    print("  headings carrying ANY direct formatting: %d  (3: must be 0)"
          % direct_on_headings)

    z = zipfile.ZipFile(str(OUT))
    doc_xml = z.read("word/document.xml").decode("utf-8")

    fonts = sorted(set(re.findall(r'w:(?:ascii|hAnsi|cs|eastAsia)="([^"]+)"', doc_xml)))
    print("  fonts named anywhere in document.xml: %s" % (fonts or "none -- all inherited"))
    bad = [f for f in fonts if any(m in f.lower() for m in MONOSPACE_MARKERS)]
    print("  monospace fonts emitted (2e: must be none): %s" % (bad or "none"))
    if bad:
        raise SystemExit("ABORT: monospace font reached the output: %s" % bad)

    core = z.read("docProps/core.xml").decode("utf-8")
    print("  core.xml     %s" % core.split("?>")[-1].strip())
    for probe in ("Mustafa", "Furat"):
        print("  %r present in core.xml: %s" % (probe, probe in core))
    print("  dc:creator element present: %s" % ("dc:creator" in core))

    styles_xml = z.read("word/styles.xml").decode("utf-8")
    root = etree.fromstring(z.read("word/styles.xml"))
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    n_styles = len(root.findall(W + "style"))
    print("  styles.xml   %d style definitions carried over" % n_styles)

    sect = etree.fromstring(z.read("word/document.xml")).find(W + "body/" + W + "sectPr")
    pgsz = sect.find(W + "pgSz")
    pgmar = sect.find(W + "pgMar")
    tw = lambda v: round(int(v) / 567.0, 2)  # twips -> cm
    print("  page size    %s x %s cm" % (tw(pgsz.get(W + "w")), tw(pgsz.get(W + "h"))))
    print("  margins      top=%s right=%s bottom=%s left=%s cm" % (
        tw(pgmar.get(W + "top")), tw(pgmar.get(W + "right")),
        tw(pgmar.get(W + "bottom")), tw(pgmar.get(W + "left"))))

    tbl_styles = sorted(set(re.findall(r'<w:tblStyle w:val="([^"]+)"', doc_xml)))
    print("  table styles referenced by the body: %s" % tbl_styles)


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    build()
